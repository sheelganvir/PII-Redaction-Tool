import os
import sys
import docx
from typing import Dict, List, Tuple, Any

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

try:
    from src.redactor import PIIRedactor
except ModuleNotFoundError:
    from redactor import PIIRedactor


class DocxRedactor:
    """
    Parses and redacts Word documents (.docx) while preserving formatting,
    runs, and table cell layout.
    """

    def __init__(self, redactor: PIIRedactor = None):
        self.redactor = redactor or PIIRedactor()

    def _redact_paragraph(self, paragraph: docx.text.paragraph.Paragraph) -> List[Dict[str, Any]]:
        """
        Redacts text in-place within individual docx runs to perfectly preserve
        run-level styling (bold, italic, font face, color, alignment, and line breaks).
        """
        full_text = paragraph.text
        if not full_text or not full_text.strip():
            return []

        # 1. Detect PII entities in full paragraph text
        entities = self.redactor.detect_entities(full_text)
        if not entities:
            return []

        # 2. Check if all detected entities exist within single runs
        all_in_single_runs = True
        if paragraph.runs:
            for ent in entities:
                orig_str = ent["text"]
                if not any(orig_str in run.text for run in paragraph.runs):
                    all_in_single_runs = False
                    break

        if all_in_single_runs and paragraph.runs:
            for ent in entities:
                orig_str = ent["text"]
                fake_str = self.redactor._get_fake_replacement(orig_str, ent["type"])
                for run in paragraph.runs:
                    if orig_str in run.text:
                        run.text = run.text.replace(orig_str, fake_str)
        else:
            # Entity spans across multiple runs: update first run with full redacted text and clear remainder
            redacted_text, _ = self.redactor.redact_text(full_text)
            if paragraph.runs:
                paragraph.runs[0].text = redacted_text
                for r in paragraph.runs[1:]:
                    r.text = ""
            else:
                paragraph.text = redacted_text

        return entities

    def redact_document(self, input_filepath: str, output_filepath: str, progress_callback=None) -> Dict[str, Any]:
        """
        Processes an entire .docx file (paragraphs and tables) and saves redacted output with optional real-time progress callbacks.
        """
        doc = docx.Document(input_filepath)
        all_detected_entities: List[Dict[str, Any]] = []

        total_paragraphs = len(doc.paragraphs)
        table_paragraphs = sum(len(cell.paragraphs) for t in doc.tables for r in t.rows for cell in r.cells)
        total_items = max(total_paragraphs + table_paragraphs, 1)
        processed_count = 0

        # 1. Redact regular paragraphs
        for idx, p in enumerate(doc.paragraphs):
            detected = self._redact_paragraph(p)
            all_detected_entities.extend(detected)
            processed_count += 1

            if progress_callback and (processed_count % 5 == 0 or processed_count == total_items):
                pct = int(15 + (processed_count / total_items) * 75)
                progress_callback(pct, f"Scanning & redacting section {processed_count} of {total_items}...")

        # 2. Redact tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        detected = self._redact_paragraph(p)
                        all_detected_entities.extend(detected)
                        processed_count += 1

                        if progress_callback and (processed_count % 10 == 0 or processed_count == total_items):
                            pct = int(15 + (processed_count / total_items) * 75)
                            progress_callback(pct, f"Processing document tables ({processed_count}/{total_items})...")

        if progress_callback:
            progress_callback(95, "Finalizing OpenXML formatting & saving redacted document...")

        doc.save(output_filepath)

        # Generate Summary Stats
        stats: Dict[str, int] = {}
        for ent in all_detected_entities:
            ent_type = ent["type"]
            stats[ent_type] = stats.get(ent_type, 0) + 1

        return {
            "total_entities_redacted": len(all_detected_entities),
            "stats_by_type": stats,
            "entities": all_detected_entities
        }
